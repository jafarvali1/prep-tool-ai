"""
Template Loader — reads the candidate's intro template (DOCX) and case study
templates (PDFs) from backend/templates/. Provides text content to AI prompts.

PDFs are image-based so we serve them as downloadable files and use their
structured text descriptions in prompts.
"""
import os
import pathlib

TEMPLATES_DIR = pathlib.Path(__file__).parent.parent / "templates"

# ─── Intro Template (from DOCX) ───────────────────────────────────────────────

# Full intro template extracted from "Candidate Interview Introduction.docx"
INTRO_TEMPLATE_TEXT = """
Candidate Interview Introduction

Hi, I am [Candidate Name].
I started my career as a software engineer. Later I moved into machine learning.
I started with MLOps. Now I work more on Generative AI and Agent based AI.

Recently, I worked on an end to end customer care AI assistant platform.
First, I did a proof of concept of RAG customer assistant.
After that, I worked on Phase one. Phase one is live in production now.

Phase one is a RAG system.
It has two main parts: Ingestion pipeline and Query pipeline.
We worked with unstructured data.
I did data cleaning and preprocessing using Docling.
For chunking and embeddings, I used Docling and sentence transformers.
I used LangChain for orchestration, prompt-templating, chaining, and Q&A.
I implemented a hybrid retriever: Semantic + Keyword.
I used Milvus as the vector database.
For generation, I used AWS Bedrock with Claude Sonnet.
I also worked on retriever evaluation and generation evaluation.
The full stack includes React for frontend, FastAPI for backend, and AWS EKS for deployment.
I also implemented guardrails.
I handled conversation history, reranking, and query optimization.

After that, we worked on Phase two — a multi agent system built using LangGraph.
It includes memory management for short term and long term memory.
There is an orchestrator agent using the supervisor pattern.
There is a router for intent detection.
Tool calling is done using MCP.

In the past, I also worked on full stack development and MLOps.
I used MLflow and Databricks, deployed and served models.
I am also familiar with SageMaker.
I did fine tuning for intent classification using a RoBERTa model.
I worked on CI/CD using GitHub, Docker, Kubernetes, ArgoCD, Helm, and Terraform.
I worked on AWS environments including EKS, ECS, S3, and Lambda.
For monitoring and observability: Datadog, CloudWatch, Prometheus, and Grafana.
"""

# ─── Case Study Template Descriptions ─────────────────────────────────────────

CASE_STUDY_TEMPLATES = {
    "mlops": {
        "name": "MLOps - Fraud Detection System (XYZ Corp)",
        "filename": "case_study_mlops.pdf",
        "description": """
MLOps Fraud Detection System case study template.
Structure used:
1. Executive Summary
2. Business Problem
3. Solution Architecture (MLOps pipeline: data ingestion → feature engineering → model training → model registry → CI/CD deployment → monitoring)
4. Technologies Used (MLflow, Databricks, Kafka, AWS SageMaker, Docker, Kubernetes)
5. Model Details (feature engineering, model selection, evaluation metrics)
6. CI/CD Pipeline (GitHub Actions, ArgoCD, Helm)
7. Monitoring & Observability (Prometheus, Grafana, Evidently for drift detection)
8. Results & Impact (fraud detection rate, false positive reduction, cost savings)
9. Challenges & Solutions
10. Lessons Learned
""",
    },
    "rag": {
        "name": "RAG Customer Call Center Agent (XYZ Corp)",
        "filename": "case_study_rag.pdf",
        "description": """
RAG (Retrieval Augmented Generation) Customer Call Center Agent case study template.
Structure used:
1. Executive Summary
2. Business Problem (call center agents searching large policy documents during live calls)
3. Solution Architecture (Ingestion pipeline + Query pipeline)
4. Ingestion Pipeline (Docling, Apache Tika, AWS Textract, Confluence APIs → S3 → preprocessing → chunking → embeddings → Milvus)
5. Query Pipeline (query cleaning → intent detection → embedding → hybrid retrieval: semantic + BM25 + Neo4j → reranking → generation via AWS Bedrock / Claude)
6. Technologies Used
7. Safety & Guardrails (PII masking with Presidio, input sandboxing, content filters)
8. Evaluation (retriever evaluation, generation evaluation, Langfuse tracing)
9. Results & Impact (reduced call time, improved accuracy, consistency)
10. Challenges & Solutions
""",
    },
    "agentic": {
        "name": "Agentic AI - Multi-Agent Customer Call Center (XYZ Corp)",
        "filename": "case_study_agentic_ai.pdf",
        "description": """
Agentic AI Multi-Agent System case study template.
Structure used:
1. Executive Summary
2. Business Problem (moving from read-only RAG to action-taking agents)
3. System Architecture (hierarchical multi-agent system with LangGraph)
4. Orchestrator Agent (ReAct reasoning, intent classification, session memory, shared scratchpad in Postgres)
5. Domain Agents:
   - Claims Agent (ReAct + RAG, tool calling to Claims Management System)
   - Billing & Accounts Agent (payment processing, human approval for high-risk)
   - Appointment & Scheduling Agent (cyclic planning, EHR APIs)
   - Policy & Benefits Agent (RAG + knowledge graph)
6. Safety Agents (Compliance Sentinel, Anti-Fraud Agent)
7. Memory Architecture (3 layers: working memory, conversation summaries, long-term vector memory)
8. Communication Pattern (Shared Blackboard via Postgres, MCP for tool calling)
9. Safety & Control (prompt injection sandboxing, tiered autonomy, audit trails)
10. Evaluation & Testing (red teaming, mock systems, task completion metrics)
11. Results & Impact
""",
    },
}


def get_intro_template() -> str:
    """Return the full intro template text from the DOCX file."""
    # Try reading from actual DOCX first
    docx_path = TEMPLATES_DIR / "intro_template.docx"
    if docx_path.exists():
        try:
            import docx
            doc = docx.Document(str(docx_path))
            texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text.strip())
            if texts:
                return "\n".join(texts)
        except Exception:
            pass
    # Fallback to hardcoded extracted text
    return INTRO_TEMPLATE_TEXT.strip()


def get_case_study_template(template_key: str = "mlops") -> dict:
    """
    Return template info for the given key ('mlops', 'rag', 'agentic').
    Returns dict with: name, filename, description, file_path (if exists)
    """
    template = CASE_STUDY_TEMPLATES.get(template_key, CASE_STUDY_TEMPLATES["mlops"])
    result = {**template}
    file_path = TEMPLATES_DIR / template["filename"]
    result["file_path"] = str(file_path) if file_path.exists() else None
    return result


def list_case_study_templates() -> list:
    """List all available case study templates."""
    return [
        {
            "key": k,
            "name": v["name"],
            "filename": v["filename"],
            "available": (TEMPLATES_DIR / v["filename"]).exists(),
        }
        for k, v in CASE_STUDY_TEMPLATES.items()
    ]


def get_template_file_path(filename: str) -> str | None:
    """Return absolute path to a template file if it exists."""
    path = TEMPLATES_DIR / filename
    return str(path) if path.exists() else None
